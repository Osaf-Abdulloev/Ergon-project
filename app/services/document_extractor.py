import io
import os
import re
import logging
from typing import Tuple, Optional

logger = logging.getLogger(__name__)

class DocumentExtractorService:
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean extracted text while preserving technical symbols, bullet points, and structure."""
        if not text:
            return ""
        
        pdf_markers = [
            r'<<.*?>>', r'/Length\s+\d+', r'/Filter\s*/FlateDecode',
            r'/MediaBox\s*\[.*?\]', r'/Parent\s+\d+\s+\d+\s+R', r'/Type\s*/\w+',
            r'/Font\s*<<.*?>>', r'%PDF-\d\.\d', r'endobj', r'stream', r'endstream'
        ]
        cleaned = text
        for pattern in pdf_markers:
            cleaned = re.sub(pattern, '', cleaned, flags=re.DOTALL | re.IGNORECASE)
            
        lines = []
        for line in cleaned.splitlines():
            line_str = line.strip()
            if not line_str or '<</Length' in line_str or '/FlateDecode' in line_str or 'Filter/FlateDecode' in line_str:
                continue
            lines.append(line_str)
            
        return "\n".join(lines).strip()

    @classmethod
    def extract_text(cls, file_bytes: bytes, filename: str, mime_type: str = "") -> Tuple[str, str]:
        """
        Extract text from file bytes using PyMuPDF (PDF), python-docx (DOCX), or OCR fallback.
        Returns tuple (extracted_text, extraction_method).
        """
        ext = os.path.splitext(filename)[1].lower()
        text = ""
        method = "unknown"

        if ext == ".pdf" or "pdf" in mime_type.lower():
            text, method = cls._extract_pdf(file_bytes)
        elif ext == ".docx" or "vnd.openxmlformats-officedocument.wordprocessingml.document" in mime_type.lower():
            text, method = cls._extract_docx(file_bytes)
        elif ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
            text, method = cls._extract_ocr_image(file_bytes)
        elif ext in [".txt", ".text", ".md", ".log"] or "text/plain" in mime_type.lower():
            text, method = cls._extract_txt(file_bytes)
        elif ext == ".doc" or "msword" in mime_type.lower():
            text, method = cls._extract_docx(file_bytes)
            if not text:
                text, method = cls._extract_txt(file_bytes)
        else:
            text, method = cls._extract_txt(file_bytes)

        # Check if text is empty or insufficient (scanned document); trigger OCR fallback
        if len(text.strip()) < 50 and ext in [".pdf", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
            logger.info("Direct text extraction yielded < 50 chars. Running OCR fallback...")
            ocr_text, ocr_method = cls._extract_ocr_pdf(file_bytes) if ext == ".pdf" else cls._extract_ocr_image(file_bytes)
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
                method = ocr_method

        return cls.clean_text(text), method

    @classmethod
    def _extract_pdf(cls, file_bytes: bytes) -> Tuple[str, str]:
        # 1. Try PyMuPDF (fitz)
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_parts = []
            for page in doc:
                page_text = page.get_text("text")
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
            extracted = "\n\n".join(text_parts)
            if len(extracted.strip()) >= 50:
                return extracted, "pymupdf"
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")

        # 2. Fallback to pypdf
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
            extracted = "\n\n".join(text_parts)
            if len(extracted.strip()) >= 50:
                return extracted, "pypdf"
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")

        return "", "pdf_failed"

    @classmethod
    def _extract_docx(cls, file_bytes: bytes) -> Tuple[str, str]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            lines = []
            for p in doc.paragraphs:
                t = p.text.strip()
                if not t:
                    continue
                if p.style and p.style.name and p.style.name.startswith("Heading"):
                    lines.append(f"\n[{p.style.name.upper()}] {t}")
                else:
                    lines.append(t)

            for table in doc.tables:
                lines.append("\n[TABLE]")
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        lines.append(" | ".join(row_text))

            extracted = "\n".join(lines)
            return extracted, "python_docx"
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}")
            return "", "docx_failed"

    @classmethod
    def _extract_txt(cls, file_bytes: bytes) -> Tuple[str, str]:
        for enc in ["utf-8", "utf-8-sig", "windows-1251", "cp1251", "latin-1"]:
            try:
                t = file_bytes.decode(enc)
                if t and t.strip():
                    return t.strip(), "txt"
            except UnicodeDecodeError:
                continue
        return "", "txt_failed"

    @classmethod
    def _extract_ocr_image(cls, file_bytes: bytes) -> Tuple[str, str]:
        # 1. Try pytesseract if available
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(img, lang="rus+eng+tgk")
            if text and text.strip():
                return text.strip(), "tesseract_ocr"
        except Exception as e:
            logger.warning(f"pytesseract image OCR failed/not configured: {e}")

        # 2. Try PaddleOCR if available
        try:
            from paddleocr import PaddleOCR
            ocr = PaddleOCR(use_angle_cls=True, lang="ru", show_log=False)
            result = ocr.ocr(file_bytes, cls=True)
            lines = []
            if result and result[0]:
                for line in result[0]:
                    lines.append(line[1][0])
            text = "\n".join(lines)
            if text and text.strip():
                return text.strip(), "paddle_ocr"
        except Exception as e:
            logger.warning(f"PaddleOCR image OCR failed/not configured: {e}")

        return "", "ocr_failed"

    @classmethod
    def _extract_ocr_pdf(cls, file_bytes: bytes) -> Tuple[str, str]:
        """Convert PDF pages into images and apply OCR."""
        try:
            import fitz
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            ocr_text_parts = []
            for page in doc:
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                p_text, _ = cls._extract_ocr_image(img_bytes)
                if p_text.strip():
                    ocr_text_parts.append(p_text.strip())
            if ocr_text_parts:
                return "\n\n".join(ocr_text_parts), "pymupdf_ocr"
        except Exception as e:
            logger.warning(f"PyMuPDF page rendering for OCR failed: {e}")

        return "", "pdf_ocr_failed"
