import io
import os
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

class CVParserService:
    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        Strips PDF stream headers and binary artifacts while preserving ALL text, bullet points, and technical symbols.
        """
        if not text:
            return ""
        
        import re
        pdf_markers = [
            r'<<.*?>>', r'/Length\s+\d+', r'/Filter\s*/FlateDecode',
            r'/MediaBox\s*\[.*?\]', r'/Parent\s+\d+\s+\d+\s+R', r'/Type\s*/\w+',
            r'/Font\s*<<.*?>>', r'%PDF-\d\.\d', r'endobj', r'stream', r'endstream'
        ]
        cleaned = text
        for pattern in pdf_markers:
            cleaned = re.sub(pattern, '', cleaned, flags=re.DOTALL | re.IGNORECASE)
            
        lines = []
        for line in cleaned.splitlines():
            line_str = line.strip()
            if not line_str or '<</Length' in line_str or '/FlateDecode' in line_str or 'Filter/FlateDecode' in line_str:
                continue
            
            # Keep all lines with printable characters, letters, digits, symbols, or bullet points
            lines.append(line_str)
            
        return "\n".join(lines).strip()


    @staticmethod
    def extract_text_from_file_bytes(file_bytes: bytes, filename: str, mime_type: str = "") -> str:
        """
        Extract clean text content from PDF, DOCX, DOC, or TXT file bytes.
        """
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf" or "pdf" in mime_type:
            return CVParserService._extract_from_pdf(file_bytes)
        elif ext == ".docx" or "vnd.openxmlformats-officedocument.wordprocessingml.document" in mime_type:
            return CVParserService._extract_from_docx(file_bytes)
        elif ext in [".txt", ".text", ".md", ".log"] or "text/plain" in mime_type:
            return CVParserService._extract_from_txt(file_bytes)
        elif ext == ".doc" or "msword" in mime_type:
            try:
                return CVParserService._extract_from_docx(file_bytes)
            except Exception:
                return CVParserService._extract_from_txt(file_bytes)
        else:
            return CVParserService._extract_from_txt(file_bytes)

    @staticmethod
    def _extract_from_pdf(file_bytes: bytes) -> str:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for idx, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(page_text.strip())
            extracted = "\n\n".join(text_parts)
            cleaned = CVParserService.clean_extracted_text(extracted)
            if cleaned.strip():
                return cleaned.strip()
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")

        # Fallback to pdfplumber if available
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                parts = [page.extract_text() for page in pdf.pages if page.extract_text()]
                extracted = "\n\n".join(parts)
                cleaned = CVParserService.clean_extracted_text(extracted)
                if cleaned.strip():
                    return cleaned.strip()
        except Exception:
            pass

        # NEVER fall back to raw binary decoding for PDFs!
        return ""

    @staticmethod
    def _extract_from_docx(file_bytes: bytes) -> str:
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        paragraphs.append(" | ".join(row_cells))

            extracted = "\n".join(paragraphs)
            cleaned = CVParserService.clean_extracted_text(extracted)
            if cleaned.strip():
                return cleaned.strip()
            raise ValueError("Empty DOCX document")
        except Exception as e:
            logger.warning(f"python-docx extraction failed: {e}")
            return ""

    @staticmethod
    def _extract_from_txt(file_bytes: bytes) -> str:
        for encoding in ["utf-8", "utf-8-sig", "windows-1251", "cp1251"]:
            try:
                text = file_bytes.decode(encoding)
                cleaned = CVParserService.clean_extracted_text(text)
                if cleaned.strip():
                    return cleaned.strip()
            except UnicodeDecodeError:
                continue
        return ""

